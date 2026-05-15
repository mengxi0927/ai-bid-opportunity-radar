from pathlib import Path

from docx import Document


DOCS = [
    Path("/Users/mengxi_yang/Downloads/招投标商机智能雷达.docx"),
    Path("/Users/mengxi_yang/Downloads/POC主要功能点.docx"),
    Path("/Users/mengxi_yang/Downloads/PRD：招投标商机智能雷达.docx"),
]


def iter_table_rows(table):
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
        if any(cells):
            yield " | ".join(cells)


def extract(path: Path) -> str:
    doc = Document(path)
    lines = [f"# {path.name}", ""]

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style else ""
            prefix = f"[{style}] " if style and style != "Normal" else ""
            lines.append(prefix + text)

    if doc.tables:
        lines.append("")
        lines.append("## Tables")
        for index, table in enumerate(doc.tables, start=1):
            lines.append(f"### Table {index}")
            lines.extend(iter_table_rows(table))
            lines.append("")

    return "\n".join(lines)


def main():
    out = Path("source-docs-extracted.md")
    out.write_text("\n\n---\n\n".join(extract(path) for path in DOCS), encoding="utf-8")
    print(out.resolve())


if __name__ == "__main__":
    main()
